from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def export_exam_docx(header, questions):
    doc = Document()

    # Header
    p_school = doc.add_paragraph()
    run = p_school.add_run(header.get("school", "").upper())
    run.bold = True
    run.font.size = Pt(12)
    p_school.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title = doc.add_paragraph()
    run = title.add_run(f"ĐỀ KIỂM TRA {header.get('semester', '').upper()} — {header.get('subject', '').upper()}")
    run.bold = True
    run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.add_run(f"{header.get('grade', '')} • Thời gian: {header.get('time', '')}").font.size = Pt(12)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if header.get("note"):
        note = doc.add_paragraph()
        note.add_run(header["note"]).font.size = Pt(11)

    doc.add_paragraph(" ")

    # Questions
    for i, q in enumerate(questions, start=1):
        p = doc.add_paragraph()
        run = p.add_run(f"Câu {i} ({q.get('points', 0)} điểm) — {q['type']}: ")
        run.bold = True
        p.add_run(q["prompt"])

        if q["type"] == "MCQ" and q.get("options"):
            for idx, opt in enumerate(q["options"]):
                if opt:
                    op = doc.add_paragraph()
                    op.add_run(f"{chr(65+idx)}. {opt}")
        elif q["type"] == "TrueFalse":
            doc.add_paragraph("Khoanh tròn Đúng hoặc Sai.")
        elif q["type"] == "FillBlank":
            doc.add_paragraph("Điền vào chỗ trống.")
        elif q["type"] == "Matching":
            doc.add_paragraph("Ghép cột A với cột B.")
        elif q["type"] == "Essay":
            doc.add_paragraph("Trình bày lời giải rõ ràng.")

    doc.add_paragraph(" ")
    ftr = doc.add_paragraph()
    ftr.add_run("— Hết —").italic = True
    ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()