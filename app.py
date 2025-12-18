import streamlit as st
import google.generativeai as genai
import pandas as pd
from io import BytesIO
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
import json
import re
from pypdf import PdfReader

# --- CẤU HÌNH TRANG ---
st.set_page_config(
    page_title="AI Exam Generator Pro (CT GDPT 2018)",
    layout="wide",
    page_icon="🏫"
)

st.title("🏫 Hệ Thống Ra Đề Thi Tiểu Học (Chuẩn CT GDPT 2018)")
st.caption("Hỗ trợ: Excel, PDF, Word. Tự động xử lý dạng câu hỏi: Đúng/Sai, Nối cột, Trắc nghiệm.")
st.markdown("---")

# ==============================================================================
# PHẦN 1: XỬ LÝ API & MODEL (CORE 7H.PY UPDATE)
# ==============================================================================

def get_best_model(api_key, response_json=False):
    """
    Tự động tìm model tốt nhất trong tài khoản.
    - Phân tích ma trận (JSON) -> Ưu tiên Flash (nhanh, context dài).
    - Viết đề (Text) -> Ưu tiên Pro (thông minh, văn hay).
    """
    genai.configure(api_key=api_key)
    try:
        all_models = list(genai.list_models())
    except Exception as e:
        return None, f"Lỗi kết nối API: {str(e)}"

    valid_models = [m.name for m in all_models if 'generateContent' in m.supported_generation_methods]
    if not valid_models:
        return None, "Không tìm thấy model nào hỗ trợ generateContent."

    # Chiến thuật chọn model
    priority = []
    if response_json:
        # Ưu tiên Flash cho JSON
        priority = [m for m in valid_models if 'flash' in m and '1.5' in m] + \
                   [m for m in valid_models if 'pro' in m and '1.5' in m]
    else:
        # Ưu tiên Pro cho Viết đề
        priority = [m for m in valid_models if 'pro' in m and '1.5' in m] + \
                   [m for m in valid_models if 'flash' in m and '1.5' in m]
    
    # Thêm các model còn lại (dự phòng)
    for m in valid_models:
        if m not in priority: priority.append(m)
        
    return priority, None

def generate_content_robust(api_key, prompt, response_json=False):
    """Hàm gọi API có cơ chế Retry (Thử lại) khi lỗi 429"""
    models, error = get_best_model(api_key, response_json)
    if error: return None, error

    last_error = ""
    config = {"response_mime_type": "application/json"} if response_json else {}

    # Thử tối đa 3 lần xoay vòng
    for attempt in range(3):
        for model_name in models:
            try:
                model = genai.GenerativeModel(model_name, generation_config=config)
                response = model.generate_content(prompt)
                return response.text, model_name
            except Exception as e:
                err_str = str(e)
                last_error = err_str
                # Nếu lỗi Quá tải (429) hoặc Model quá tải (503)
                if "429" in err_str or "ResourceExhausted" in err_str or "503" in err_str:
                    time.sleep(2) # Nghỉ 2s rồi thử model khác
                    continue
                continue 

    return None, f"Thất bại sau nhiều lần thử. Lỗi cuối: {last_error}"

# ==============================================================================
# PHẦN 2: BỘ ĐỌC FILE ĐA NĂNG (PRE-PROCESSORS)
# ==============================================================================

def process_excel_to_text(file):
    try:
        # Đọc không header để bắt trọn dữ liệu
        df = pd.read_excel(file, header=None)
        
        # Tìm dòng Header chính
        header_idx = 0
        for idx, row in df.iterrows():
            row_str = row.astype(str).str.lower().values
            if any('chủ đề' in s or 'mạch' in s for s in row_str):
                header_idx = idx
                break
        
        df_clean = df.iloc[header_idx:].reset_index(drop=True)
        
        # QUAN TRỌNG: Forward Fill để xử lý Merge Cell (File Book1.xlsx của bạn bị lỗi này)
        # Các ô chủ đề bị gộp sẽ được điền tên xuống các dòng dưới
        df_clean = df_clean.ffill()
        
        return df_clean.to_string()
    except Exception as e:
        return f"Lỗi đọc Excel: {e}"

def process_pdf_to_text(file):
    try:
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        return f"Lỗi đọc PDF: {e}"

def process_docx_to_text(file):
    try:
        doc = docx.Document(file)
        text = ""
        for table in doc.tables:
            for row in table.rows:
                # Dùng dấu | để ngăn cách các cột cho AI dễ hiểu
                text += " | ".join([cell.text.strip() for cell in row.cells]) + "\n"
        return text
    except Exception as e:
        return f"Lỗi đọc Word: {e}"

# ==============================================================================
# PHẦN 3: LOGIC AI (PHÂN TÍCH & TẠO ĐỀ)
# ==============================================================================

def analyze_matrix_step(file_text, api_key):
    """Bước 1: Chuyển văn bản thô thành cấu trúc JSON"""
    prompt = f"""
    Bạn là trợ lý xử lý dữ liệu giáo dục. Hãy phân tích văn bản ma trận đề thi dưới đây thành JSON.
    
    VĂN BẢN ĐẦU VÀO:
    {file_text[:20000]} 

    YÊU CẦU OUTPUT (JSON List):
    Hãy trích xuất danh sách các yêu cầu ra đề. Chỉ lấy những dòng có số lượng câu hỏi > 0.
    Cấu trúc mẫu:
    [
      {{
        "topic": "Tên chủ đề / Mạch kiến thức",
        "yccd": "Yêu cầu cần đạt (nếu có)",
        "questions": [
           {{"type": "TN nhiều lựa chọn", "level": "Biết", "count": "1 câu"}},
           {{"type": "TN Đúng/Sai", "level": "Hiểu", "count": "1 câu"}},
           {{"type": "Tự luận", "level": "Vận dụng", "count": "1 câu"}}
        ]
      }}
    ]
    Lưu ý:
    - Nếu gặp "Đúng - Sai" hãy ghi type là "TN Đúng/Sai".
    - Nếu gặp "Nối cột" hãy ghi type là "TN Nối cột".
    """
    res, model = generate_content_robust(api_key, prompt, response_json=True)
    return res, model

def create_exam_step(blueprint_json, subject, api_key):
    """Bước 2: Viết đề thi (Strict Mode - Ép đúng dạng bài)"""
    prompt = f"""
    Bạn là giáo viên tiểu học (CT GDPT 2018). Hãy soạn đề thi môn {subject} dựa trên cấu trúc JSON sau:
    {blueprint_json}

    QUY TẮC BẮT BUỘC VỀ DẠNG CÂU HỎI (STRICT MODE):
    1. Dạng "TN nhiều lựa chọn":
       - Hỏi 1 câu, có 4 đáp án A, B, C, D.
    
    2. Dạng "TN Đúng/Sai" (Bắt buộc làm đúng format này):
       - Đưa ra 1 câu dẫn chính.
       - Bên dưới là 4 ý a), b), c), d).
       - Học sinh sẽ xác định mỗi ý là Đúng hay Sai.
       - Ví dụ:
         Câu 1: Phát biểu nào sau đây về...
         a) ... (Đ/S?)
         b) ... (Đ/S?)
    
    3. Dạng "TN Nối cột" (Matching):
       - Tạo Cột A (1, 2, 3, 4) và Cột B (a, b, c, d).
       - Yêu cầu nối thông tin tương ứng.

    4. Dạng "Tự luận":
       - Câu hỏi mở, ngắn gọn, sát thực tế.

    CẤU TRÚC ĐỀ THI:
    - PHẦN I: TRẮC NGHIỆM (Bao gồm nhiều lựa chọn, đúng/sai, nối cột).
    - PHẦN II: TỰ LUẬN.
    - PHẦN III: ĐÁP ÁN VÀ HƯỚNG DẪN CHẤM (Chi tiết thang điểm).

    TRÌNH BÀY:
    - Đánh số câu liên tục (Câu 1, Câu 2...).
    - Ngôn ngữ trong sáng, dễ hiểu.
    """
    res, model = generate_content_robust(api_key, prompt, response_json=False)
    return res, model

# ==============================================================================
# PHẦN 4: XUẤT FILE WORD ĐẸP
# ==============================================================================

def create_word_doc(text):
    doc = docx.Document()
    
    # Cài đặt Font chữ toàn bài
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    
    # Căn lề A4 chuẩn
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    lines = text.split('\n')
    for line in lines:
        clean_line = line.strip()
        if not clean_line: continue
        
        p = doc.add_paragraph(clean_line)
        lower = clean_line.lower()
        
        # Logic in đậm thông minh
        # 1. In đậm Tiêu đề lớn (Phần I, Phần II, Đề thi...)
        if any(x in lower for x in ["phần i", "phần ii", "phần iii", "đề thi", "đáp án", "hướng dẫn chấm"]):
            runner = p.runs[0]
            runner.bold = True
            runner.font.size = Pt(14)
            runner.font.color.rgb = RGBColor(0, 0, 0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        # 2. In đậm đầu câu hỏi (Câu 1:, Câu 2:...)
        # Regex check: Bắt đầu bằng "Câu" + số + dấu chấm hoặc hai chấm
        elif re.match(r'^Câu\s+\d+[:.]', clean_line):
            p.runs[0].bold = True
            
        # 3. In đậm các ý a), b) trong câu Đúng/Sai nếu cần (Tùy chọn)
        
    bio = BytesIO()
    doc.save(bio)
    return bio

# ==============================================================================
# PHẦN 5: GIAO DIỆN STREAMLIT (UI)
# ==============================================================================

with st.sidebar:
    st.header("🔑 Cấu hình")
    api_key = st.text_input("Nhập Google API Key", type="password")
    st.info("Hệ thống tự động chọn model tốt nhất (Flash/Pro) để tránh lỗi.")

col1, col2 = st.columns([1, 1.5])

with col1:
    st.subheader("1. Thiết lập")
    uploaded_file = st.file_uploader("Upload Ma Trận (Excel, PDF, Word)", type=['xlsx', 'pdf', 'docx'])
    subject_name = st.text_input("Tên môn & Lớp (VD: Khoa học lớp 4)")
    
    # Nút thực hiện 2 bước
    if st.button("🚀 Phân tích & Tạo đề", type="primary"):
        if not uploaded_file or not api_key or not subject_name:
            st.warning("Vui lòng nhập đủ: API Key, File và Tên môn.")
        else:
            status = st.status("Đang xử lý...", expanded=True)
            try:
                # --- BƯỚC 1: ĐỌC FILE ---
                status.write("📂 Đang đọc nội dung file...")
                file_text = ""
                if uploaded_file.name.endswith('.xlsx'):
                    file_text = process_excel_to_text(uploaded_file)
                elif uploaded_file.name.endswith('.pdf'):
                    file_text = process_pdf_to_text(uploaded_file)
                else:
                    file_text = process_docx_to_text(uploaded_file)
                
                # --- BƯỚC 2: PHÂN TÍCH ---
                status.write("🤖 Đang phân tích ma trận (Trích xuất JSON)...")
                blueprint, m1 = analyze_matrix_step(file_text, api_key)
                
                if blueprint:
                    st.session_state['blueprint'] = blueprint
                    # Clean json string nếu AI trả về format markdown ```json ... ```
                    clean_bp = blueprint.replace("```json", "").replace("```", "").strip()
                    
                    status.write(f"✅ Phân tích xong (Model: {m1})")
                    
                    # --- BƯỚC 3: TẠO ĐỀ ---
                    status.write("✍️ Đang viết đề (Strict Mode - Đúng dạng bài)...")
                    exam_txt, m2 = create_exam_step(clean_bp, subject_name, api_key)
                    
                    if exam_txt:
                        st.session_state['exam_result'] = exam_txt
                        status.update(label=f"Hoàn tất! (Model: {m2})", state="complete", expanded=False)
                    else:
                        status.update(label="Lỗi tạo đề", state="error")
                        st.error(m2)
                else:
                    status.update(label="Lỗi phân tích", state="error")
                    st.error(m1)
                    
            except Exception as e:
                status.update(label="Lỗi hệ thống", state="error")
                st.error(str(e))

with col2:
    st.subheader("2. Kết quả")
    
    tab1, tab2 = st.tabs(["📝 Đề thi hoàn chỉnh", "🔍 Cấu trúc phân tích (Debug)"])
    
    with tab2:
        if 'blueprint' in st.session_state:
            st.caption("Đây là những gì AI đọc được từ file của bạn:")
            try:
                # Cố gắng parse JSON để hiển thị đẹp
                bp_json = st.session_state['blueprint'].replace("```json", "").replace("```", "").strip()
                st.json(json.loads(bp_json))
            except:
                st.text(st.session_state['blueprint'])
        else:
            st.info("Chưa có dữ liệu.")

    with tab1:
        if 'exam_result' in st.session_state:
            # Cho phép sửa trực tiếp
            edited_text = st.text_area("Xem và sửa đề trước khi tải:", 
                                     value=st.session_state['exam_result'], 
                                     height=700)
            
            # Tạo file word
            doc_file = create_word_doc(edited_text)
            
            st.download_button(
                label="📥 Tải xuống file Word (.docx)",
                data=doc_file.getvalue(),
                file_name=f"De_thi_{subject_name.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary"
            )
        else:
            st.info("Kết quả sẽ hiển thị tại đây.")
